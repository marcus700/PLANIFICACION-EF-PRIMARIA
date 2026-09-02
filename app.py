import io
import re
import time
import base64
import urllib.parse
import requests
from PIL import Image
import docx
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from google import genai
from google.genai import types
import streamlit as st

# ==============================================================================
# IMPORTACIÓN DESDE TU BASE DE DATOS EXTERNA (cneb_datos.py)
# ==============================================================================
from cneb_datos import CNEB_PRIMARIA, obtener_ciclo_primaria

# DICCIONARIO DE CAPACIDADES OFICIALES DEL CNEB - EDUCACIÓN FÍSICA
CAPACIDADES_CNEB = {
    "Se desenvuelve de manera autónoma a través de su motricidad": [
        "Comprende su cuerpo.",
        "Se expresa corporalmente."
    ],
    "Asume una vida saludable": [
        "Comprende las relaciones entre la actividad física, alimentación, postura e higiene personal y del ambiente, y la salud.",
        "Incorpora prácticas que mejoran su calidad de vida."
    ],
    "Interactúa a través de sus habilidades sociomotrices": [
        "Se relaciona utilizando sus habilidades sociomotrices.",
        "Crea y aplica estrategias y tácticas de juego."
    ]
}

def normalizar_grado_cneb(grado_str: str) -> str:
    """Mapea la opción seleccionada al formato de llave exacta en cneb_datos.py"""
    if "1" in grado_str: return "1° de Primaria"
    if "2" in grado_str: return "2° de Primaria"
    if "3" in grado_str: return "3° de Primaria"
    if "4" in grado_str: return "4° de Primaria"
    if "5" in grado_str: return "5° de Primaria"
    if "6" in grado_str: return "6° de Primaria"
    return "2° de Primaria"

# ==============================================================================
# CONFIGURACIÓN DE PÁGINA Y CSS MEJORADO PARA ALTA VISIBILIDAD DE BOTONES
# ==============================================================================
st.set_page_config(
    page_title="PlanificaEF Primaria - Plataforma de Educación Física",
    page_icon="⚽",
    layout="wide"
)

st.markdown("""
<style>
    header, footer, [data-testid="stHeader"], [data-testid="stDecoration"], 
    [data-testid="stStatusWidget"], [data-testid="stViewerBadge"], 
    [data-testid="manage-app-button"], .stAppDeployButton, .viewerBadge_container__1613n {
        display: none !important;
        visibility: hidden !important;
        opacity: 0 !important;
    }
    
    html, body, [data-testid="stAppViewContainer"], .stApp {
        background-color: #F8FAFC !important;
        color: #0F172A !important;
    }
    
    .block-container {
        padding-top: 1.2rem !important;
        padding-bottom: 1rem !important;
    }
    
    .main-header {
        font-size: 2.2rem;
        color: #1E3A8A;
        text-align: center;
        font-weight: 800;
        margin-bottom: 0.2rem;
    }
    .sub-header {
        font-size: 1.05rem;
        color: #4B5563;
        text-align: center;
        margin-bottom: 1.5rem;
    }

    .stTextInput input, .stTextArea textarea, .stSelectbox [data-baseweb="select"] {
        background-color: #FFFFFF !important;
        color: #0F172A !important;
        border: 1px solid #94A3B8 !important;
        border-radius: 8px !important;
    }

    /* ESTILOS DE ALTO CONTRASTE Y LEGIBILIDAD PARA BOTONES DE HERRAMIENTAS */
    div.st-key-btn_unidad > button {
        background: linear-gradient(135deg, #7C3AED 0%, #5B21B6 100%) !important;
        background-color: #7C3AED !important;
        border: 2px solid #6D28D9 !important;
        border-radius: 12px !important;
        box-shadow: 0 4px 14px rgba(124, 58, 237, 0.4) !important;
        transition: all 0.3s ease !important;
    }
    div.st-key-btn_unidad > button:hover {
        transform: translateY(-2px) !important;
        box-shadow: 0 6px 18px rgba(124, 58, 237, 0.6) !important;
    }

    div.st-key-btn_proyecto > button {
        background: linear-gradient(135deg, #059669 0%, #047857 100%) !important;
        background-color: #059669 !important;
        border: 2px solid #047857 !important;
        border-radius: 12px !important;
        box-shadow: 0 4px 14px rgba(5, 150, 105, 0.4) !important;
        transition: all 0.3s ease !important;
    }
    div.st-key-btn_proyecto > button:hover {
        transform: translateY(-2px) !important;
        box-shadow: 0 6px 18px rgba(5, 150, 105, 0.6) !important;
    }

    div.st-key-btn_sesion > button {
        background: linear-gradient(135deg, #2563EB 0%, #1D4ED8 100%) !important;
        background-color: #2563EB !important;
        border: 2px solid #1D4ED8 !important;
        border-radius: 12px !important;
        box-shadow: 0 4px 14px rgba(37, 99, 235, 0.4) !important;
        transition: all 0.3s ease !important;
    }
    div.st-key-btn_sesion > button:hover {
        transform: translateY(-2px) !important;
        box-shadow: 0 6px 18px rgba(37, 99, 235, 0.6) !important;
    }

    div.stButton > button:not([key="btn_unidad"]):not([key="btn_proyecto"]):not([key="btn_sesion"]) {
        background: linear-gradient(135deg, #1E40AF 0%, #1E3A8A 100%) !important;
        background-color: #1E40AF !important;
        border: 2px solid #1D4ED8 !important;
        border-radius: 12px !important;
        padding: 0.75rem 1.5rem !important;
        box-shadow: 0 4px 16px rgba(30, 64, 175, 0.5) !important;
        transition: all 0.3s ease !important;
    }
    div.stButton > button:not([key="btn_unidad"]):not([key="btn_proyecto"]):not([key="btn_sesion"]):hover {
        background: linear-gradient(135deg, #2563EB 0%, #1E40AF 100%) !important;
        transform: translateY(-2px) !important;
        box-shadow: 0 6px 20px rgba(37, 99, 235, 0.7) !important;
    }

    div.stButton > button,
    div.stButton > button *,
    div.stButton > button p,
    div.stButton > button span,
    div.stButton > button div {
        color: #FFFFFF !important;
        font-weight: 800 !important;
        font-size: 1.05rem !important;
        text-shadow: 0px 1px 3px rgba(0, 0, 0, 0.6) !important;
    }
</style>
""", unsafe_allow_html=True)

st.markdown('<div class="main-header">⚽ PlanificaEF - Plataforma de Educación Física</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-header">Sistema de Planificación Curricular Especializado en Educación Física Primaria (CNEB - MINEDU)</div>', unsafe_allow_html=True)

# ==============================================================================
# CONTROL DE ACCESO
# ==============================================================================
def check_password():
    if "password_correct" not in st.session_state:
        st.session_state["password_correct"] = False

    if st.session_state["password_correct"]:
        return True

    st.markdown("---")
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.subheader("🔒 Acceso Restringido")
        pwd_input = st.text_input("Contraseña de acceso:", type="password", key="pwd_input")
        if st.button("Ingresar 🚀"):
            target_pwd = st.secrets.get("APP_PASSWORD", "docente2026ef")
            if pwd_input == target_pwd:
                st.session_state["password_correct"] = True
                st.rerun()
            else:
                st.error("❌ Contraseña incorrecta.")
    return False

if not check_password():
    st.stop()

# MEMORIA PERSISTENTE
if 'resultado_md' not in st.session_state:
    st.session_state['resultado_md'] = None
if 'tipo_doc_generado' not in st.session_state:
    st.session_state['tipo_doc_generado'] = None
if 'fname_clean' not in st.session_state:
    st.session_state['fname_clean'] = None
if 'tipo_documento' not in st.session_state:
    st.session_state['tipo_documento'] = "Unidad de Aprendizaje"
if 'imagenes_dict' not in st.session_state:
    st.session_state['imagenes_dict'] = {}

# SIDEBAR CON MODELOS ESTABLES DE GOOGLE STUDIO Y OPENAI (OPCIONAL)
st.sidebar.title("⚙️ Configuración EF")
if st.sidebar.button("🔒 Cerrar Sesión"):
    st.session_state["password_correct"] = False
    st.rerun()

st.sidebar.markdown("---")
if "GEMINI_API_KEY" in st.secrets and st.secrets["GEMINI_API_KEY"]:
    api_key = st.secrets["GEMINI_API_KEY"]
    st.sidebar.success("🔑 Gemini API Key activada.")
else:
    api_key = st.sidebar.text_input("🔑 Google AI Studio API Key:", type="password")

if "OPENAI_API_KEY" in st.secrets and st.secrets["OPENAI_API_KEY"]:
    openai_api_key = st.secrets["OPENAI_API_KEY"]
    st.sidebar.success("🎨 OpenAI Key activada.")
else:
    openai_api_key = st.sidebar.text_input("🎨 OpenAI API Key (Opcional):", type="password")

# OPCIONES DE MODELOS OFICIALES Y ESTABLES
model_choice = st.sidebar.selectbox(
    "Modelo de Gemini:", 
    ["gemini-3.6-flash", "gemini-3.5-flash", "gemini-2.5-pro", "gemini-2.5-flash", "gemini-2.0-flash", "gemini-1.5-flash", "gemini-1.5-pro"]
)

# ==============================================================================
# HERRAMIENTAS DE EDUCACIÓN FÍSICA
# ==============================================================================
st.markdown("### 📋 Selecciona el Documento de Educación Física a Elaborar:")

col_b1, col_b2, col_b3 = st.columns(3)

with col_b1:
    if st.button("📘 Unidad de Aprendizaje", key="btn_unidad", use_container_width=True):
        st.session_state['tipo_documento'] = "Unidad de Aprendizaje"
        st.rerun()

with col_b2:
    if st.button("🚀 Proyecto de Aprendizaje", key="btn_proyecto", use_container_width=True):
        st.session_state['tipo_documento'] = "Proyecto de Aprendizaje"
        st.rerun()

with col_b3:
    if st.button("🏃 Sesión de Aprendizaje de Ed. Física", key="btn_sesion", use_container_width=True):
        st.session_state['tipo_documento'] = "Sesión de Aprendizaje de Ed. Física"
        st.rerun()

tipo_documento = st.session_state['tipo_documento']

COLOR_MAP = {
    "Unidad de Aprendizaje": "#7C3AED",
    "Proyecto de Aprendizaje": "#059669",
    "Sesión de Aprendizaje de Ed. Física": "#2563EB"
}
banner_color = COLOR_MAP.get(tipo_documento, "#7C3AED")

st.markdown(f"""
<div style="background-color: {banner_color}; color: white; padding: 0.6rem 1rem; border-radius: 8px; font-weight: bold; font-size: 1.1rem; margin-top: 0.8rem; margin-bottom: 1.2rem; text-align: center; text-shadow: 0px 1px 3px rgba(0,0,0,0.4);">
    📍 Área Exclusiva: EDUCACIÓN FÍSICA | Herramienta: {tipo_documento.upper()}
</div>
""", unsafe_allow_html=True)

# ==============================================================================
# MOTOR DE ILUSTRACIONES INFOGRÁFICAS TIPO FICHA MINEDU PERÚ
# ==============================================================================
def construir_prompt_minedu(tipo_actividad, desc_actividad):
    """Crea la estructura del prompt para imitar la lámina infográfica educativa del MINEDU"""
    base = (
        "MINEDU Peru primary school Physical Education textbook infographic sheet. "
        "Official Peruvian curriculum 2D vector storybook style, white background card with clean border. "
        "Peruvian primary school children wearing PE sports uniforms (white t-shirt, blue athletic shorts, sneakers). "
        "Pedagogical exercise diagrams with cones, hoops, sports balls, clear dashed movement arrows, step numbers (1, 2, 3), "
        "and clear instructional visual panels. "
    )
    
    if "1." in tipo_actividad or "Activación" in tipo_actividad:
        return (
            f"{base} Blue themed header banner: '1. ACTIVACIÓN FISIOLÓGICA'. "
            f"Top panel: Physical education teacher with whistle and smiling Peruvian kids running in groups around orange cones, "
            f"with an inset diagram box for 'MOVILIDAD ARTICULAR' showing neck, shoulders, hips and ankle rotations. "
            f"Bottom panel: Kids doing skipping, lunges, and heel kicks, with a wrist pulse checking cartoon and a 15-second stopwatch icon. "
            f"{desc_actividad}."
        )
    elif "2." in tipo_actividad or "Básica" in tipo_actividad:
        return (
            f"{base} Green themed header banner: '2. ACTIVIDAD BÁSICA (FAMILIARIZACIÓN Y EXPLORACIÓN MOTRIZ)'. "
            f"Top panel: 'Pases con obstáculos' with two Peruvian students passing a ball through an upright hoop stand with cones and distance markers '5m'. "
            f"Bottom panel: 'Mini-fútbol de pases' showing small court, cone goalposts, dashed passing arrows, and teamwork banner. "
            f"{desc_actividad}."
        )
    elif "3." in tipo_actividad or "Avanzada" in tipo_actividad:
        return (
            f"{base} Orange themed header banner: '3. ACTIVIDAD AVANZADA (PROGRESIÓN PEDAGÓGICA Y COMPLEJIZACIÓN)'. "
            f"Top panel: 'Circuito de precisión y velocidad' with step numbers (1) Zigzag between orange cones, (2) Jumping inside 3 consecutive floor hoops, (3) Throwing into an elevated hanging hoop target. "
            f"Bottom panel: Tactical divided court diagram with children playing in two teams. "
            f"{desc_actividad}."
        )
    else:
        return (
            f"{base} Magenta/Pink themed header banner: '4. ACTIVIDAD DE APLICACIÓN (JUEGO MODIFICADO)'. "
            f"Main panel: Isometric court diagram with two pitches 'CAMPO A' and 'CAMPO B', small goals, Peruvian school children playing in teams, rules panel, rotation arrows, "
            f"and a bottom inset box for 'PAUSA DE HIDRATACIÓN' showing a student drinking from a water bottle with a 2-minute timer icon. "
            f"{desc_actividad}."
        )

def generar_imagen_actividad_universal(openai_key, tipo_actividad, prompt_actividad):
    """Genera la lámina infográfica didáctica estilo MINEDU Perú de forma inmediata"""
    prompt_completo = construir_prompt_minedu(tipo_actividad, prompt_actividad)
    
    # 1. Si el usuario ingresó clave de OpenAI, intenta generar con DALL-E 3
    if openai_key and len(openai_key.strip()) > 10:
        try:
            headers = {
                "Authorization": f"Bearer {openai_key.strip()}",
                "Content-Type": "application/json"
            }
            payload = {
                "model": "dall-e-3",
                "prompt": prompt_completo,
                "n": 1,
                "size": "1024x1024",
                "quality": "standard"
            }
            resp = requests.post("https://api.openai.com/v1/images/generations", headers=headers, json=payload, timeout=35)
            data = resp.json()
            if "data" in data and len(data["data"]) > 0:
                img_url = data["data"][0]["url"]
                resp_img = requests.get(img_url, timeout=25)
                if resp_img.status_code == 200:
                    return Image.open(io.BytesIO(resp_img.content)), None
        except Exception:
            pass

    # 2. Generador Directo e Ilimitado (Estilo Ficha MINEDU Perú)
    try:
        encoded_prompt = urllib.parse.quote(prompt_completo)
        url = f"https://image.pollinations.ai/prompt/{encoded_prompt}?width=1024&height=768&nologo=true"
        resp = requests.get(url, timeout=30)
        if resp.status_code == 200:
            return Image.open(io.BytesIO(resp.content)), None
    except Exception as e:
        return None, str(e)

    return None, "Error al generar la infografía MINEDU."

# ==============================================================================
# CONVERTIDOR DE MARKDOWN A WORD CON TABLAS EN TONOS PASTELES
# ==============================================================================
def add_formatted_text(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        else:
            paragraph.add_run(part)

def markdown_to_docx(md_text, ie_nombre="I.E. N° 22314", es_horizontal=False):
    doc = docx.Document()
    PASTEL_COLORS = ['D9E1F2', 'E2EFDA', 'FFF2CC', 'E8D8F8', 'E0F2FE', 'FCE4D6']
    table_count = 0
    
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
        if es_horizontal:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Inches(11.69)
            section.page_height = Inches(8.27)
        else:
            section.orientation = WD_ORIENT.PORTRAIT
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)

    lines = md_text.split('\n')
    in_table = False
    table_data = []

    def render_table(t_data, color_hex):
        rows = len(t_data)
        cols = max(len(r) for r in t_data) if rows > 0 else 0
        if rows > 0 and cols > 0:
            t = doc.add_table(rows=rows, cols=cols)
            t.style = 'Table Grid'
            for r_idx, row_cells in enumerate(t_data):
                for c_idx, cell_value in enumerate(row_cells):
                    if c_idx < cols:
                        cell = t.cell(r_idx, c_idx)
                        p_cell = cell.paragraphs[0]
                        p_cell.text = ""
                        add_formatted_text(p_cell, cell_value)
                        
                        if r_idx == 0:
                            shading_elm = OxmlElement('w:shd')
                            shading_elm.set(qn('w:val'), 'clear')
                            shading_elm.set(qn('w:color'), 'auto')
                            shading_elm.set(qn('w:fill'), color_hex)
                            cell._tc.get_or_add_tcPr().append(shading_elm)
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.color.rgb = RGBColor(30, 58, 138)
                                    run.font.bold = True

    for line in lines:
        line_str = line.strip()
        line_str = re.sub(r'<br\s*/?>', ' ', line_str)
        
        if line_str.startswith('|') and line_str.endswith('|'):
            in_table = True
            if re.match(r'^\|[\s\:\-\|]+\|$', line_str):
                continue
            cells = [c.strip() for c in line_str.split('|')[1:-1]]
            table_data.append(cells)
            continue
        elif in_table:
            if table_data:
                table_count += 1
                header_color = PASTEL_COLORS[(table_count - 1) % len(PASTEL_COLORS)]
                render_table(table_data, header_color)
            in_table = False
            table_data = []

        heading_match = re.match(r'^(#{1,6})\s*(.*)$', line_str)
        if heading_match:
            hashes = heading_match.group(1)
            title_text = heading_match.group(2).strip()
            level = len(hashes)
            
            p = doc.add_paragraph()
            if level in [1, 2]:
                run = p.add_run(title_text.replace('**', ''))
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(30, 58, 138)
            elif level in [3, 4]:
                run = p.add_run(title_text.replace('**', ''))
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(30, 58, 138)
            else:
                add_formatted_text(p, title_text)
            continue

        if line_str.startswith('• ') or line_str.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            clean_bullet = line_str[2:].strip()
            add_formatted_text(p, clean_bullet)
        elif line_str != "":
            p = doc.add_paragraph()
            add_formatted_text(p, line_str)

    if in_table and table_data:
        table_count += 1
        header_color = PASTEL_COLORS[(table_count - 1) % len(PASTEL_COLORS)]
        render_table(table_data, header_color)
        in_table = False
        table_data = []
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==============================================================================
# FORMULARIO DE DATOS
# ==============================================================================
st.subheader(f"📝 Configuración de Datos para Educación Física: {tipo_documento}")

c1, c2, c3 = st.columns(3)
with c1:
    dre_ugel = st.text_input("DRE / UGEL:", "Ica / Ica")
    ie_nombre = st.text_input("Institución Educativa:", "N.° 22314 'Vicenta Aquije de Huamán'")
with c2:
    director = st.text_input("Directora:", "Prof. Luisa Ruth Aronés Herrera")
    docente = st.text_input("Docente de Educación Física:", "Mario A. García Torres")
with c3:
    if tipo_documento == "Unidad de Aprendizaje":
        ciclo_seleccionado_str = st.selectbox(
            "Ciclo CNEB:", 
            ["III Ciclo (1° y 2° Grado)", "IV Ciclo (3° y 4° Grado)", "V Ciclo (5° y 6° Grado)"], 
            index=0
        )
        ciclo_actual = ciclo_seleccionado_str.split(" (")[0]

        if "III" in ciclo_actual:
            grado_1_ciclo_str = "1er Grado"
            grado_2_ciclo_str = "2do Grado"
            grado_1_cneb = "1° de Primaria"
            grado_2_cneb = "2° de Primaria"
        elif "IV" in ciclo_actual:
            grado_1_ciclo_str = "3er Grado"
            grado_2_ciclo_str = "4to Grado"
            grado_1_cneb = "3° de Primaria"
            grado_2_cneb = "4° de Primaria"
        else: # V Ciclo
            grado_1_ciclo_str = "5to Grado"
            grado_2_ciclo_str = "6to Grado"
            grado_1_cneb = "5° de Primaria"
            grado_2_cneb = "6° de Primaria"
        
        grado_seccion = f"{grado_1_ciclo_str} y {grado_2_ciclo_str}"
        st.info(f"Ciclo Seleccionado: **{ciclo_actual}**")
        grado_normalizado_cneb = None
    else:
        grado_seccion = st.selectbox("Grado y Sección:", ["1er Grado A", "2do Grado A", "3er Grado A", "4to Grado A", "5to Grado A", "6to Grado A"], index=1)
        grado_normalizado_cneb = normalizar_grado_cneb(grado_seccion)
        ciclo_actual = obtener_ciclo_primaria(grado_normalizado_cneb)
        st.info(f"Ciclo CNEB Detectado: **{ciclo_actual}**")

# VARIABLES ESPECÍFICAS PARA CADA HERRAMIENTA
if tipo_documento == "Sesión de Aprendizaje de Ed. Física":
    f1, f2, f3 = st.columns(3)
    with f1:
        num_doc = st.text_input("N.° de Sesión:", "01")
    with f2:
        fecha_sugerida = st.text_input("Fecha:", "22 de junio de 2026")
    with f3:
        duracion_sesion = st.selectbox("Duración de la Clase:", ["45 minutos", "90 minutos", "135 minutos"], index=1)
    
    st.markdown("##### 📌 Configuración Pedagógica de la Sesión")
    titulo_sesion_input = st.text_input("Título de la Actividad / Sesión de Clase (Opcional):", value="", placeholder="Ej. Leemos señales para desplazarnos y reconocer direcciones en el patio")
    
    col_s1, col_s2 = st.columns(2)
    with col_s1:
        comps_seleccionadas = st.multiselect(
            "Competencia(s) a Trabajar:",
            list(CAPACIDADES_CNEB.keys()),
            default=["Se desenvuelve de manera autónoma a través de su motricidad"]
        )
        
        capacidades_disponibles = []
        for comp in comps_seleccionadas:
            for cap in CAPACIDADES_CNEB.get(comp, []):
                capacidades_disponibles.append(f"{comp}: {cap}")
        
        capacidades_seleccionadas = st.multiselect(
            "Capacidad(es) a Utilizar (Se actualizan según la competencia):",
            options=capacidades_disponibles,
            default=capacidades_disponibles,
            help="Selecciona las capacidades que se trabajarán en esta sesión."
        )

        estandar_custom = st.text_area("Estándar de la Competencia (Opcional - Blanco para automático):", value="", height=70, placeholder="Texto del estándar...")
    with col_s2:
        tipo_motivacion = st.selectbox(
            "Tipo de Motivación para el Inicio de la Clase:",
            ["A través de una actividad física", "A través de una imagen", "A través de una historia"],
            index=0
        )
        criterios_custom = st.text_area("Criterios de Evaluación (Opcional - Blanco para automático):", value="", height=70, placeholder="Ej. 1. Ejecuta desplazamientos orientados en el patio. 2. Identifica nociones de derecha e izquierda.")
        evidencia_custom = st.text_input("Evidencia de Aprendizaje (Opcional - Blanco para automático):", value="", placeholder="Ej. Ejecución de desplazamientos coordinados hacia señales leídas.")

    # CUADRO DE MATERIALES A UTILIZAR EN LA SESIÓN
    st.markdown("##### 🎒 Cuadro de Recursos y Materiales a Utilizar en la Sesión:")
    col_mat1, col_mat2 = st.columns(2)
    with col_mat1:
        materiales_patio = st.text_area(
            "⚽ Materiales Deportivos y del Patio:",
            value="Conos, aros, balones, silbato, tiza para delimitar el patio, colchonetas.",
            height=70
        )
    with col_mat2:
        materiales_estudiante = st.text_area(
            "🧴 Recursos de Higiene y del Estudiante:",
            value="Botella de agua personal, toalla de mano, jabón, polo de cambio deportivo.",
            height=70
        )

    fechas_duracion = fecha_sugerida
    duracion_semanas = 1
    sesiones_por_semana = 1
    producto_unidad = ""
    problema_contexto = titulo_sesion_input.strip() if titulo_sesion_input.strip() else "Desarrollo de nociones espaciales, coordinación motriz y convivencia en juegos de Educación Física."

else:  # Unidad o Proyecto EF
    f1, f2, f3, f4, f5 = st.columns(5)
    with f1:
        num_doc = st.text_input("N.° de Unidad / Proyecto:", "04")
    with f2:
        fechas_duracion = st.text_input("Fechas / Periodo:", "Del 22 de junio al 17 de julio de 2026")
        fecha_sugerida = fechas_duracion
    with f3:
        duracion_semanas = st.slider("Número de Semanas:", min_value=2, max_value=8, value=4)
    with f4:
        sesiones_por_semana = st.selectbox("Sesiones por Semana:", [1, 2, 3], index=1)
    with f5:
        producto_unidad = st.text_input("Producto Final Tangible:", "Festival Lúdico-Motor Peruanito")
        duracion_sesion = "90 minutos"

    problema_contexto = st.text_area(
        "📋 Describe el Tema, Problema de Contexto o Necesidad Motriz/Saludable de los Estudiantes:",
        height=120,
        value="Dificultades de coordinación motriz, orientación espacial en el patio al desplazarse en grupo, poco conocimiento de juegos tradicionales peruanos y falta de hábitos de higiene personal (lavado de manos, cambio de polo) al finalizar la actividad física."
    )
    capacidades_seleccionadas = []
    materiales_patio = ""
    materiales_estudiante = ""

# ==============================================================================
# PROMPTS ESPECIALIZADOS QUE LEEN DE cneb_datos.py
# ==============================================================================

def generar_prompt_unidad_ef_10_secciones():
    cneb_datos_text = ""
    for comp_nombre, comp_info in CNEB_PRIMARIA.items():
        est_txt = comp_info["estandares"].get(ciclo_actual, "")
        des_list_g1 = comp_info["desempenos"].get(grado_1_cneb, [])
        des_list_g2 = comp_info["desempenos"].get(grado_2_cneb, [])
        
        cneb_datos_text += f"\n\nCOMPETENCIA: {comp_nombre}\nESTÁNDAR OFICIAL ({ciclo_actual}):\n{est_txt}"
        cneb_datos_text += f"\nDESEMPEÑOS OFICIALES ({grado_1_ciclo_str}):\n" + "\n".join(des_list_g1)
        cneb_datos_text += f"\nDESEMPEÑOS OFICIALES ({grado_2_ciclo_str}):\n" + "\n".join(des_list_g2)

    total_sesiones_unidad = duracion_semanas * sesiones_por_semana

    return f"""
Actúa como un especialista en currículo educativo peruano y docente experto en el área de Educación Física para Educación Básica Regular (CNEB). 

Tu tarea es elaborar una UNIDAD DE APRENDIZAJE completa para el CICLO COMPLETO, con una MATRIZ DE PLANIFICACIÓN y una SECUENCIA DE SESIONES diferenciada para cada grado.

🚨 REGLAS CRÍTICAS DE COMPLETITUD Y ESTRUCTURA POR CICLO (OBLIGATORIO LLEGAR HASTA LA SECCIÓN X):
1. FINALIZA EL DOCUMENTO OBLIGATORIAMENTE HASTA LA SECCIÓN X (RECURSOS Y FIRMAS).
2. EN LA SECCIÓN VIII (MATRIZ DE PLANIFICACIÓN), genera DOS PLANIFICACIONES COMPLETAS: una para {grado_1_ciclo_str} y otra para {grado_2_ciclo_str}. Dentro de cada una, agrupa las sesiones por competencia (C1, C2, C3), encabezando cada grupo con su estándar resaltando con negrita lo evaluado.
3. EN LA SECCIÓN IX (SECUENCIA DE SESIONES), genera DOS SECUENCIAS COMPLETAS: una tabla para {grado_1_ciclo_str} y otra para {grado_2_ciclo_str}.
4. Usa los DESEMPEÑOS ESPECÍFICOS del grado correspondiente para cada matriz.
5. REGLA OBLIGATORIA DE CRITERIOS: Para cada sesión individual de la matriz, debes formular OBLIGATORIAMENTE EXACTAMENTE 3 CRITERIOS DE EVALUACIÓN claros, observables y medibles. Cada criterio debe estar redactado de forma fluida e integrada (Acción + Contenido + Condición) pero SIN ESCRIBIR NI MOSTRAR VISIBLEMENTE las palabras/etiquetas 'Acción:', 'Contenido:' ni 'Condición:' (debe ser una sola oración continua por criterio).
6. Usa la fórmula para la formulacion de titulos de las  sesiones : "¡Verbo en plural + acción directa + reto físico!", Sin metáforas, sin fantasía, El lenguaje debe ser sencillo, claro y de fácil lectura (evita tecnicismos abstractos como "patrones de locomoción" o "segmentos corporales" en el título). 

DATOS OFICIALES EXTRAÍDOS DE cneb_datos.py PARA ESTA UNIDAD ({ciclo_actual}):
{cneb_datos_text}

DATOS PARA LA GENERACIÓN:
- N° de Unidad: Unidad N° {num_doc}
- Ciclo / Grados: {ciclo_actual} - ({grado_1_ciclo_str} y {grado_2_ciclo_str})
- Nombre de la IE: {ie_nombre}
- Nombre del Docente: {docente}
- Nombre del Director(a): {director}
- Duración / Fechas: {duracion_semanas} semanas ({total_sesiones_unidad} sesiones en total) - ({fechas_duracion})
- Tema central / Problemática a abordar: {problema_contexto}
- Producto de la Unidad: {producto_unidad}

---

ESTRUCTURA OBLIGATORIA DE LA UNIDAD DE APRENDIZAJE:

(Desarrolla las Secciones I a VII de manera general para el ciclo)
1. TÍTULO DE LA UNIDAD
2. II. DATOS INFORMATIVOS
3. III. SITUACIÓN SIGNIFICATIVA
- Contextualizar la realidad motriz y de salud de los estudiantes relacionada con la problemática: {problema_contexto}.
- Plantear 3 preguntas retadoras/desafiantes asociadas a la solución motriz.
- Proponer la estrategia pedagógica para resolver el reto (circuitos, festivales lúdico-motores, juegos tradicionales, etc.).
4. IV. PRODUCTO DE LA UNIDAD
- Describir un desempeño práctico o un producto tangible/demostrable claro: {producto_unidad}.
5. V. ENFOQUES TRANSVERSALES
- Seleccionar 2 enfoques transversales del CNEB.
- Especificar en tabla 3 columnas Enfoque Transversal, Valor(es) y Acciones o Actitudes Observables adaptadas a Educación Física.
6. VI. COMPETENCIAS TRANSVERSALES
- Incluir en tabla 3 columnas "Gestiona su aprendizaje de manera autónoma" y "Se desenvuelve en entornos virtuales generados por las TIC" con sus respectivas Capacidades y Desempeños aplicados al área.
7. VII. ESTÁNDARES, COMPETENCIAS Y CAPACIDADES DEL ÁREA DE EDUCACIÓN FÍSICA
- Transcribir las 3 competencias oficiales del área con sus capacidades y estándares completos del ciclo correspondiente ({ciclo_actual}):
  * Competencia 1: Se desenvuelve de manera autónoma a través de su motricidad.
  * Competencia 2: Asume una vida saludable.
  * Competencia 3: Interactúa a través de sus habilidades sociomotrices.

8. VIII. MATRIZ DE PLANIFICACIÓN POR CICLO (ORGANIZADA POR GRADO Y COMPETENCIA)
A continuación, desarrolla las DOS planificaciones completas.

---
### **MATRIZ DE PLANIFICACIÓN PARA {grado_1_ciclo_str.upper()}**
(Usa los desempeños de {grado_1_cneb})

**COMPETENCIA 1: Se desenvuelve de manera autónoma a través de su motricidad**
> **ESTÁNDAR CNEB COMPLETO ({ciclo_actual}):** [Transcribe aquí el estándar COMPLETO de la Competencia 1]
| Sesiones de C1 | Competencia / Capacidad | Desempeño precisado completo (con **negrita**) (de {grado_1_cneb}) | Criterios de Evaluación | Evidencia | Instrumento |

**COMPETENCIA 2: Asume una vida saludable**
> **ESTÁNDAR CNEB COMPLETO ({ciclo_actual}):** [Transcribe aquí el estándar COMPLETO de la Competencia 2]
| Sesiones de C2 | Competencia / Capacidad | Desempeño precisado completo (con **negrita**) (de {grado_1_cneb}) | Criterios de Evaluación | Evidencia | Instrumento |

**COMPETENCIA 3: Interactúa a través de sus habilidades sociomotrices**
> **ESTÁNDAR CNEB COMPLETO ({ciclo_actual}):** [Transcribe aquí el estándar COMPLETO de la Competencia 3]
| Sesiones de C3 | Competencia / Capacidad | Desempeño precisado completo (con **negrita**) (de {grado_1_cneb}) | Criterios de Evaluación | Evidencia | Instrumento |

---
### **MATRIZ DE PLANIFICACIÓN PARA {grado_2_ciclo_str.upper()}**
(Usa los desempeños de {grado_2_cneb})

**COMPETENCIA 1: Se desenvuelve de manera autónoma a través de su motricidad**
> **ESTÁNDAR CNEB COMPLETO ({ciclo_actual}):** [Transcribe aquí el estándar COMPLETO de la Competencia 1]
| Sesiones de C1 | Competencia / Capacidad | Desempeño precisado completo (con **negrita**) (de {grado_2_cneb}) | Criterios de Evaluación | Evidencia | Instrumento |

**COMPETENCIA 2: Asume una vida saludable**
> **ESTÁNDAR CNEB COMPLETO ({ciclo_actual}):** [Transcribe aquí el estándar COMPLETO de la Competencia 2]
| Sesiones de C2 | Competencia / Capacidad | Desempeño precisado completo (con **negrita**) (de {grado_2_cneb}) | Criterios de Evaluación | Evidencia | Instrumento |

**COMPETENCIA 3: Interactúa a través de sus habilidades sociomotrices**
> **ESTÁNDAR CNEB COMPLETO ({ciclo_actual}):** [Transcribe aquí el estándar COMPLETO de la Competencia 3]
| Sesiones de C3 | Competencia / Capacidad | Desempeño precisado completo (con **negrita**) (de {grado_2_cneb}) | Criterios de Evaluación | Evidencia | Instrumento |
---

9. IX. SECUENCIA DE SESIONES (Diferenciada por Grado)
Genera DOS TABLAS DE SECUENCIA DE SESIONES completas e independientes, una para cada grado del ciclo.

---
### **SECUENCIA DE SESIONES PARA {grado_1_ciclo_str.upper()}**
(Tabla completa con las {total_sesiones_unidad} sesiones para el primer grado del ciclo)
| N° | Título de la actividad | Propósito de la actividad | Representación gráfica |

---
### **SECUENCIA DE SESIONES PARA {grado_2_ciclo_str.upper()}**
(Tabla completa con las {total_sesiones_unidad} sesiones, adaptando los títulos y propósitos para el segundo grado del ciclo para reflejar mayor complejidad)
| N° | Título de la actividad (adaptado) | Propósito de la actividad (adaptado) | Representación gráfica |
---

10. X. RECURSOS
- Recursos para el Docente, Recursos para el Estudiante.
- Fecha y espacio para firmas.
"""

def generar_prompt_proyecto_ef():
    cneb_datos_text = ""
    for comp_nombre, comp_info in CNEB_PRIMARIA.items():
        est_txt = comp_info["estandares"].get(ciclo_actual, "")
        des_list = comp_info["desempenos"].get(grado_normalizado_cneb, [])
        cneb_datos_text += f"\n\nCOMPETENCIA: {comp_nombre}\nESTÁNDAR OFICIAL ({ciclo_actual}):\n{est_txt}\nDESEMPEÑOS OFICIALES ({grado_normalizado_cneb}):\n" + "\n".join(des_list)

    total_sesiones_proyecto = duracion_semanas * sesiones_por_semana

    return f"""
Actúa como un Especialista Pedagógico experto en Educación Física del Ministerio de Educación de Perú (MINEDU). Tu tarea es diseñar un Proyecto de Aprendizaje completo bajo el enfoque por competencias del Currículo Nacional de la Educación Básica (CNEB), manteniendo de manera estricta y detallada una estructura formal sin cortar el documento al final.

🚨 REGLAS CRÍTICAS DE COMPLETITUD Y SÍNTESIS EN TABLAS (OBLIGATORIO LLEGAR HASTA LA SECCIÓN IX):
1. DEBES DESARROLLAR EL PROYECTO COMPLETO LLEGANDO OBLIGATORIAMENTE HASTA LA SECCIÓN IX (RECURSOS Y MATERIALES Y FIRMAS DE DIRECTORA Y DOCENTE). QUEDA STRICTAMENTE PROHIBIDO CORTAR EL DOCUMENTO.
2. MANTÉN LAS RESPUESTAS Y TEXTOS DENTRO DE LAS CELDAS DE LAS TABLAS DE FORMA SINTÉTICA Y CONCISA (1 A 2 LÍNEAS POR CELDA) PARA GARANTIZAR QUE EL DOCUMENTO SE GENERE COMPLETO.
3. EN LA SECCIÓN VII (CUADRO CRONOLÓGICO DE SESIONES), DESARROLLA LAS {total_sesiones_proyecto} SESIONES ({duracion_semanas} semanas, {sesiones_por_semana} sesión(es) por semana) UNA POR UNA EN LA TABLA DE 3 COLUMNAS. ESTÁ PROHIBIDO USAR PUNTOS SUSPENSIVOS (...) O OMITIR SESIONES.
4. EN LA SECCIÓN VI (MATRIZ DE PROPÓSITOS), ORGANIZA LA MATRIZ COMPETENCIA POR COMPETENCIA (C1, C2, C3). Para cada competencia, coloca arriba su Estándar COMPLETO con negrita en lo utilizado, y en la tabla transcribe el Desempeño COMPLETO con **negrita** en lo movilizado y precisado, incluyendo OBLIGATORIAMENTE EXACTAMENTE 3 CRITERIOS DE EVALUACIÓN por cada sesión.
5. REGLA OBLIGATORIA DE CRITERIOS: Cada criterio debe estar redactado de forma fluida e integrada (Acción + Contenido + Condición) pero SIN ESCRIBIR NI MOSTRAR VISIBLEMENTE las palabras/etiquetas 'Acción:', 'Contenido:' ni 'Condición:' (debe ser una sola oración continua por criterio).

DATOS OFICIALES EXTRAÍDOS DE cneb_datos.py PARA ESTE PROYECTO ({grado_seccion} - {ciclo_actual}):
{cneb_datos_text}

DATOS PARA LA GENERACIÓN:
- N° de Proyecto: Proyecto N° {num_doc}
- DRE / UGEL: {dre_ugel}
- Institución Educativa: {ie_nombre}
- Nivel: Educación Primaria
- Ciclo: {ciclo_actual}
- Grado y Sección: {grado_seccion}
- Área Curricular: Educación Física
- Duración y Frecuencia: {duracion_semanas} semanas, {sesiones_por_semana} sesiones por semana = {total_sesiones_proyecto} sesiones en total ({fechas_duracion})
- Tema o Problemática Central: {problema_contexto}
- Producto Final: {producto_unidad}

---

ESTRUCTURA OBLIGATORIA DEL PROYECTO DE APRENDIZAJE DE EDUCACIÓN FÍSICA:

I. TÍTULO DEL PROYECTO
- Debe ser motivador, creativo, retador y entre comillas (Ejemplo: "¡CELEBRAMOS NUESTRA PERUANIDAD EN EL GRAN FESTIVAL LÚDICO-MOTOR!").

II. DATOS INFORMATIVOS
- DRE/UGEL, IE, Nivel, Ciclo, Grado y Sección, Área (Educación Física), Duración, N° de sesiones, Docente, Director(a).

III. SITUACIÓN SIGNIFICATIVA
Redacta una situación basada en un contexto real de la escuela en 4 bloques sintéticos:
- Contexto, Problema o necesidad ({problema_contexto}), Reto (2 a 3 preguntas) y Propósito.

IV. CUADRO DE ENFOQUES TRANSVERSALES
Elabora una tabla con 1 o 2 enfoques transversales más pertinentes (Enfoque, Valores, Actitudes observables).

V. CUADRO DE NEGOCIACIÓN Y PLANIFICACIÓN CON LOS ESTUDIANTES
Tabla sintética de 4 columnas (¿Qué queremos hacer?, ¿Cómo lo haremos?, ¿Qué necesitamos?, ¿Cómo nos daremos cuenta de que lo logramos?) con respuestas realistas de asamblea.

VI. CUADRO DE PROPÓSITOS DE APRENDIZAJE Y EVALUACIÓN MATRIZADA (ORGANIZADO COMPETENCIA POR COMPETENCIA)
Desarrolla 3 bloques independientes, UNO POR CADA COMPETENCIA DE EDUCACIÓN FÍSICA:

1. **COMPETENCIA 1: Se desenvuelve de manera autónoma a través de su motricidad**
   > **ESTÁNDAR CNEB COMPLETO ({ciclo_actual}):** [Texto íntegro del estándar del ciclo transcrito literalmente sin recortar, con **negrita** en la parte movilizada]
   
   Tabla de la Competencia 1 (con las sesiones asociadas a esta competencia):
   | Sesión / Actividad | Desempeño CNEB Completo (con **negrita**) | EXACTAMENTE 3 Criterios de Evaluación por Sesión (Integrados sin etiquetas) | Evidencia de Aprendizaje | Instrumento de Evaluación |

2. **COMPETENCIA 2: Asume una vida saludable**
   > **ESTÁNDAR CNEB COMPLETO ({ciclo_actual}):** [Texto íntegro del estándar del ciclo transcrito literalmente sin recortar, con **negrita** en la parte movilizada]
   
   Tabla de la Competencia 2 (con las sesiones asociadas a esta competencia):
   | Sesión / Actividad | Desempeño CNEB Completo (con **negrita**) | EXACTAMENTE 3 Criterios de Evaluación por Sesión (Integrados sin etiquetas) | Evidencia de Aprendizaje | Instrumento de Evaluación |

3. **COMPETENCIA 3: Interactúa a través de sus habilidades sociomotrices**
   > **ESTÁNDAR CNEB COMPLETO ({ciclo_actual}):** [Texto íntegro del estándar del ciclo transcrito literalmente sin recortar, con **negrita** en la parte movilizada]
   
   Tabla de la Competencia 3 (con las sesiones asociadas a esta competencia):
   | Sesión / Actividad | Desempeño CNEB Completo (con **negrita**) | EXACTAMENTE 3 Criterios de Evaluación por Sesión (Integrados sin etiquetas) | Evidencia de Aprendizaje | Instrumento de Evaluación |

*REGLA DEL DESEMPEÑO EN CADA TABLA: Copia el desempeño oficial completo del CNEB para {grado_seccion}, RESALTANDO EN NEGRITA lo utilizado y la precisión/contexto agregado.*

VII. PLANIFICACIÓN CRONOLÓGICA DETALLADA DE LAS SESIONES
Desglosa secuencialmente las {total_sesiones_proyecto} sesiones. Tabla obligatoria de 3 COLUMNAS:
| Denominación de la sesión | Propósito detallado de la sesión | Representación gráfica |
- Denominación: Número y título motivador entre comillas.
- Propósito detallado: Síntesis clara que incluya calentamiento, juego práctico e higiene personal.
- Representación gráfica: Descripción breve de la imagen o esquema del patio.

VIII. PRODUCTOS DEL PROYECTO
- Producto Intangible / Práctico (Festival, Mini olimpiadas, Gincana).
- Producto Tangible ({producto_unidad}).

IX. RECURSOS Y MATERIALES
- Material deportivo, material reciclado, materiales de higiene.
- Espacios educativos y espacio para firmas de la Directora y Docente de Educación Física.
"""

def generar_prompt_sesion_ef():
    comps_str = ", ".join(comps_seleccionadas) if comps_seleccionadas else "Seleccionar automáticamente según el tema del CNEB"
    
    if capacidades_seleccionadas:
        cap_str = "\n".join([f"- {c}" for c in capacidades_seleccionadas])
    else:
        cap_str = "Generar automáticamente según la(s) competencia(s) elegida(s)"

    est_str = estandar_custom.strip() if estandar_custom.strip() else "Transcribir el Estándar COMPLETO oficial del ciclo del CNEB con negrita en la parte movilizada"
    crit_str = criterios_custom.strip() if criterios_custom.strip() else "Formular automáticamente mínimo 3 criterios claros con la estructura Acción + Contenido + Condición"
    evid_str = evidencia_custom.strip() if evidencia_custom.strip() else "Generar automáticamente la evidencia motriz o demostración práctica adecuada"
    
    mat_patio_str = materiales_patio.strip() if materiales_patio.strip() else "Conos, aros, balones, silbato, colchonetas."
    mat_est_str = materiales_estudiante.strip() if materiales_estudiante.strip() else "Botella de agua personal, toalla pequeña, jabón, polo de cambio."

    return f"""
Actúa como Docente Experto en Educación Física para Primaria bajo el enfoque oficial del CNEB del MINEDU Perú.
Elabora una SESIÓN DE CLASE PRÁCTICA DE EDUCACIÓN FÍSICA completa para {grado_seccion} ({ciclo_actual}).

DATOS INGRESADOS PARA LA SESIÓN:
- N.° de Sesión: {num_doc}
- Título de la actividad: "{problema_contexto}"
- IE: {ie_nombre} | Docente: {docente} | Fecha: {fecha_sugerida} | Duración: {duracion_sesion}
- Tipo de Motivación elegida: {tipo_motivacion}
- Competencia(s) solicitada(s): {comps_str}
- Capacidades solicitadas:
{cap_str}
- Estándar solicitado: {est_str}
- Criterios solicitados: {crit_str}
- Evidencia solicitada: {evid_str}
- Materiales Deportivos y del Patio: {mat_patio_str}
- Recursos de Higiene y del Estudiante: {mat_est_str}

---

REGLAS DE FORMATO Y ESTRUCTURA OBLIGATORIA DE LA SESIÓN (DESARROLLAR COMPLETA DE PRINCIPIO A FIN):

1. ENCABEZADO Y TÍTULO DE LA SESIÓN:
Muestra EXACTAMENTE la siguiente estructura en la parte superior:
# **SESIÓN DE APRENDIZAJE DE EDUCACIÓN FÍSICA N.º {num_doc}**
## **"{problema_contexto.upper()}"**
*(QUEDA STRICTAMENTE PROHIBIDO COLOCAR CUALQUIER OTRO DATO, FECHA O SUBTÍTULO DEBAJO DEL TÍTULO DE LA SESIÓN).*

2. I: DATOS INFORMATIVOS
| DATOS INFORMATIVOS | DETALLE |
| Institución Educativa | {ie_nombre} |
| Docente de Educación Física | {docente} |
| Grado y Sección | {grado_seccion} ({ciclo_actual}) |
| Fecha | {fecha_sugerida} |
| Duración | {duracion_sesion} |

3. II: PROPÓSITOS DE APRENDIZAJE Y EVIDENCIAS
> **ESTÁNDAR CNEB COMPLETO ({ciclo_actual}):** [Texto íntegro del estándar del ciclo con **negrita** en la parte aplicada]

| ÁREA | COMPETENCIA Y CAPACIDADES | DESEMPEÑO PRECISADO COMPLETO (con **negrita**) | EXACTAMENTE 3 CRITERIOS DE EVALUACIÓN (Integrados sin etiquetas) | PROPÓSITO DE LA CLASE | EVIDENCIA | INSTRUMENTO |
- **Capacidades a incluir:** Incluye textualmente las capacidades indicadas:
{cap_str}
- **Criterios de Evaluación:** Redacta OBLIGATORIAMENTE EXACTAMENTE 3 CRITERIOS DE EVALUACIÓN claros, observables y medibles que integren de forma fluida e implícita los tres elementos pedagógicos (**Acción + Contenido + Condición**), pero QUEDA STRICTAMENTE PROHIBIDO escribir o visualizar las palabras/etiquetas 'Acción:', 'Contenido:' o 'Condición:' en el texto (debe ser una sola oración continua y natural por criterio).

4. III: ENFOQUE TRANSVERSAL (ÚNICO Y ESPECÍFICO)
| ENFOQUE TRANSVERSAL PRIORIZADO | VALOR(ES) | ACTITUDES OBSERVABLES |

5. IV: COMPETENCIAS TRANSVERSALES
| COMPETENCIA TRANSVERSAL | CAPACIDADES | DESEMPEÑOS PRECISADOS |

6. V: PREPARACIÓN DE LA CLASE Y MATERIALES A UTILIZAR
Elabora un cuadro detallado considerando los siguientes materiales indicados:
| ¿Qué necesitamos hacer antes de la sesión de Ed. Física? | ¿Qué recursos y materiales se utilizarán? |
- En la columna de materiales, incluye de forma organizada:
  * **Materiales deportivos y del patio:** {mat_patio_str}
  * **Materiales de higiene y del estudiante:** {mat_est_str}

7. MOMENTOS DE LA CLASE DE EDUCACIÓN FÍSICA:

- **INICIO (Aprox. ... min):**
  Redactado en PRIMERA PERSONA DEL PLURAL Y TIEMPO PRESENTE. Debe considerar ESTRICTAMENTE el siguiente orden:
  1. **Motivación ({tipo_motivacion}):** [Desarrollar la motivación según el tipo elegido: {tipo_motivacion}].
  2. **Saberes previos:** [Preguntas abiertas sobre el tema/movimientos].
  3. **Problematización / Conflicto cognitivo:** [Reto motriz o pregunta desafiante].
  4. **Propósito de la clase:** [Comunicar qué aprenderán hoy].
  5. **Criterios de evaluación:** [Explicar cómo serán evaluados].
  6. **Acuerdos de convivencia:** [2 a 3 normas de seguridad en el patio].

- **DESARROLLO (Aprox. ... min) - ALTO DETALLE DIDÁCTICO, RIGOR MOTOR Y METODOLÓGICO:**
  Redactado en PRIMERA PERSONA DEL PLURAL Y TIEMPO PRESENTE. Para CADA una de las siguientes actividades debes describir obligatoriamente:
  *Nombre de la actividad entre comillas*, *Organización y materiales en el patio*, *Descripción paso a paso y consigna pedagógica docente*, *Reglas claras / Sistema de juego* y *Variante de progresión o reto*.

  1. **Activación Fisiológica (Calentamiento dinámico y preventivo):**
     - *Movilidad Articular y Trote Lúdico:* [Describir una dinámica lúdica de desplazamiento por todo el patio con cambios de ritmo y movilidad articular progresiva cefalocaudal].
     - *Estiramientos activos y coordinación:* [Describir ejercicios dinámicos de activación muscular y toma de pulso inicial].

  2. **Actividad Básica (Familiarización y Exploración Motriz - 2 Actividades detalladas):**
     - *Actividad A (Exploración individual / en parejas):* [Nombre del juego]. Detalla la organización en el patio, el uso del material, la consigna técnica paso a paso (postura, apoyos, orientación) y cómo el docente orienta la exploración.
     - *Actividad B (Interacción socio-motriz en pequeños grupos):* [Nombre del juego]. Detalla la formación de grupos, delimitación del espacio con conos/tiza, reglas iniciales de cooperación o pases, y el objetivo lúdico a cumplir.

  3. **Actividad Avanzada (Progresión Pedagógica y Complejización - 2 Actividades detalladas):**
     - *Actividad A (Complejización motriz y encadenamiento):* [Nombre del juego]. Detalla cómo se combinan 2 o más patrones motores (ej. desplazarse con cambio de dirección + salvar obstáculos + precisión en lanzamiento), exigiendo mayor coordinación, ritmo y fluidez.
     - *Actividad B (Variabilidad, reto y oposición suave):* [Nombre del juego]. Detalla la modificación de condiciones (reducción de tiempo, límite de toques, roles de defensores/atacantes) para forzar la toma rápida de decisiones motrices.

  4. **Actividad de Aplicación (Transferencia y Consolidación en Juego Modificado - 1 Juego Principal):**
     - *Nombre del Juego de Aplicación:* [Nombre del gran juego modificado/cooperativo].
     - *Organización y espacio del patio:* [Distribución de zonas, arcos/metas, delimitaciones y distribución de los equipos].
     - *Reglas completas y sistema de puntuación:* [Explicar detalladamente cómo se anota punto, qué acciones están prohibidas por seguridad, cómo rotan los roles y cómo se fomenta la estrategia grupal].
     - *Pausa breve de hidratación consciente:* [Indicación para rehidratarse con agua antes del cierre].

- **CIERRE (Aprox. ... min) - OBLIGATORIO Y COMPLETO:**
  1. **Actividad de Recuperación (Vuelta a la calma):** Ejercicios de respiración guiada (inhalación/exhalación profunda) y estiramientos suaves con música o silencio.
  2. **Metacognición motriz:** Redacta de 3 a 4 preguntas reflexivas explícitas (¿Qué aprendimos sobre nuestro cuerpo? ¿Qué dificultades tuvimos y cómo las superamos? ¿Para qué nos sirve lo practicado hoy en nuestra vida diaria?).
  3. **Rutina Obligatoria de Higiene Personal:** Describe en detalle la práctica autónoma de aseo personal, lavado de manos con jabón, secado con toalla y cambio de polo deportivo.

8. VI: LISTA DE COTEJO DE EDUCACIÓN FÍSICA (Genera una tabla limpia con los criterios de evaluación e incluye 8 a 10 estudiantes ficticios representativos para optimizar espacio y garantizar la completitud del documento).
"""

# ==============================================================================
# EJECUCIÓN CON SISTEMA DUAL ROBUSTO ANTI-404 Y COMPLETITUD
# ==============================================================================
st.markdown("---")

if st.button(f"✨ Generar {tipo_documento}"):
    if not api_key:
        st.error("⚠️ Ingresa tu API Key de Google AI Studio en la barra lateral izquierda.")
    elif not problema_contexto:
        st.warning("⚠️ Completa el campo del Tema o Problemática de Educación Física.")
    else:
        try:
            client = genai.Client(api_key=api_key)
            st.session_state['imagenes_dict'] = {}
            
            if tipo_documento == "Unidad de Aprendizaje":
                prompt_maestro = generar_prompt_unidad_ef_10_secciones()
            elif tipo_documento == "Proyecto de Aprendizaje":
                prompt_maestro = generar_prompt_proyecto_ef()
            else:
                prompt_maestro = generar_prompt_sesion_ef()

            sys_inst = """Eres un Especialista Curricular del MINEDU Perú dedicado exclusivamente al área de Educación Física. Generas documentos pedagógicos completos en Markdown alineados estrictamente al CNEB.
REGLA ABSOLUTA DE FORMATO DE VIÑETAS: Para viñetas utiliza únicamente el guion ("- ") seguido del texto en negrita con doble asterisco ("**Título:** detalle"). Queda terminantemente PROHIBIDO usar asteriscos duplicados como "* *" o dejar asteriscos flotantes.
REGLA ABSOLUTA DE COMPLETITUD: Queda STRICTAMENTE PROHIBIDO recortar, abreviar o dejar incompletos los documentos al final.
Para evitar que el documento se corte al final, debes ser SINTÉTICO, CONCISO Y DIRECTO dentro de las celdas de las tablas, garantizando que el documento se redacte ENTERO de principio a fin, concluyendo obligatoriamente en la última sección con el espacio para firmas correspondientes."""

            with st.spinner(f"⚽ Google Gemini está redactando tu {tipo_documento} para {grado_seccion}..."):
                config = types.GenerateContentConfig(
                    system_instruction=sys_inst,
                    temperature=0.10,
                    max_output_tokens=8192
                )
                
                modelos_a_probar = [
                    model_choice,
                    "gemini-3.6-flash",
                    "gemini-3.5-flash",
                    "gemini-2.5-pro",
                    "gemini-2.5-flash",
                    "gemini-2.0-flash",
                    "gemini-1.5-flash",
                    "gemini-1.5-pro"
                ]
                
                response = None
                ultimo_err = None
                
                for mod in modelos_a_probar:
                    try:
                        response = client.models.generate_content(
                            model=mod,
                            contents=prompt_maestro,
                            config=config
                        )
                        if response and response.text:
                            break
                    except Exception as err:
                        ultimo_err = err
                        continue
                
                if not response or not response.text:
                    raise ultimo_err
                
                st.session_state['resultado_md'] = response.text
                st.session_state['tipo_doc_generado'] = tipo_documento
                st.session_state['fname_clean'] = f"{tipo_documento.replace(' ', '_')}_EF_{ciclo_actual.replace(' ', '_')}.docx"
                
                st.success(f"✅ ¡{tipo_documento} de Educación Física generado con éxito!")

        except Exception as e:
            err_str = str(e)
            if "429" in err_str or "RESOURCE_EXHAUSTED" in err_str:
                st.warning("⏳ Límite de velocidad alcanzado. Por favor, espera 60 segundos y vuelve a intentarlo.")
            else:
                st.error(f"❌ Ocurrió un error con la API de Google AI Studio: {err_str}")

# ==============================================================================
# DESPLIEGUE DE RESULTADOS Y DESCARGA EN WORD
# ==============================================================================
if st.session_state['resultado_md'] is not None:
    st.markdown("---")
    
    es_sesion = st.session_state['tipo_doc_generado'] == "Sesión de Aprendizaje de Ed. Física"
    
    if es_sesion:
        tab_preview, tab_img, tab_download = st.tabs([
            "📄 Vista Previa (Permanente)",
            "🖼️ Láminas Infográficas MINEDU",
            "📥 Descargar en Word (.docx)"
        ])
    else:
        tab_preview, tab_download = st.tabs([
            "📄 Vista Previa (Permanente)",
            "📥 Descargar en Word (.docx)"
        ])
    
    with tab_preview:
        st.markdown(st.session_state['resultado_md'])
        
    if es_sesion:
        with tab_img:
            st.markdown("### 🏃‍♂️ Láminas Infográficas Ilustradas del Desarrollo (Fichas MINEDU Perú)")
            st.info("💡 Haz clic en cada botón para generar la lámina didáctica completa con diagramas de cancha, flechas tácticas, caricaturas escolares y recuadros de pulso/hidratación.")
            
            actividades = [
                (
                    "1. ACTIVACIÓN FISIOLÓGICA (CALENTAMIENTO DINÁMICO Y PREVENTIVO)",
                    f"Calentamiento dinámico en el patio escolar: trote grupal con conos, panel de movilidad articular y toma de pulso cardíaco para {problema_contexto}"
                ),
                (
                    "2. ACTIVIDAD BÁSICA (FAMILIARIZACIÓN Y EXPLORACIÓN MOTRIZ)",
                    f"Pases con obstáculos con aros verticales en soporte con cotas de 5m y mini-fútbol de pases obligatorios con banner de trabajo en equipo para {problema_contexto}"
                ),
                (
                    "3. ACTIVIDAD AVANZADA (PROGRESIÓN PEDAGÓGICA Y COMPLEJIZACIÓN)",
                    f"Circuito de agilidad numerado 1 Zigzag con conos, 2 Saltos en 3 aros, 3 Lanzamiento a diana alta, y juego táctico de balón torre para {problema_contexto}"
                ),
                (
                    "4. ACTIVIDAD DE APLICACIÓN (TRANSFERENCIA Y JUEGO MODIFICADO)",
                    f"Losa deportiva dividida en Campo A y Campo B con arcos y juego deportivo cooperativo por equipos, rotación y pausa de hidratación consciente con botella de agua y cronómetro para {problema_contexto}"
                )
            ]
            
            cols_img = st.columns(2)
            for idx, (fase, desc) in enumerate(actividades):
                col_idx = idx % 2
                with cols_img[col_idx]:
                    st.markdown(f"#### 🏅 {fase}")
                    
                    if fase in st.session_state['imagenes_dict']:
                        item = st.session_state['imagenes_dict'][fase]
                        st.image(item["img"], caption=item["desc"], use_container_width=True)
                        
                        buf = io.BytesIO()
                        item["img"].save(buf, format="PNG")
                        st.download_button(
                            label=f"⬇️ Descargar Lámina MINEDU ({fase.split(' ')[1]})",
                            data=buf.getvalue(),
                            file_name=f"lamina_minedu_{idx+1}.png",
                            mime="image/png",
                            key=f"dl_act_{idx}"
                        )
                    else:
                        st.caption(desc)
                        if st.button(f"🎨 Generar Lámina MINEDU ({fase.split(' ')[1]})", key=f"btn_indiv_{idx}", use_container_width=True):
                            with st.spinner(f"Diseñando lámina infográfica de {fase}..."):
                                img_res, err = generar_imagen_actividad_universal(openai_api_key, fase, desc)
                                if img_res:
                                    st.session_state['imagenes_dict'][fase] = {"img": img_res, "desc": desc}
                                    st.rerun()
                                else:
                                    st.error(f"Error al generar: {err}")
                    st.markdown("---")

    with tab_download:
        es_horizontal_doc = st.session_state['tipo_doc_generado'] in ["Unidad de Aprendizaje", "Proyecto de Aprendizaje"]
        
        buffer_doc = markdown_to_docx(
            st.session_state['resultado_md'], 
            ie_nombre=ie_nombre,
            es_horizontal=es_horizontal_doc
        )
        
        st.download_button(
            label=f"💾 Descargar {st.session_state['tipo_doc_generado']} en Word (.docx)",
            data=buffer_doc,
            file_name=st.session_state['fname_clean'],
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
        st.info("💡 **Nota:** El documento Word incluye la insignia editable y las tablas en tonos pasteles.")
